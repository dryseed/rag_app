import sqlite3
from datetime import datetime, timezone
from typing import List, Optional, Dict
import re
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph


DB_PATH = 'qa_knowledge.db'

CREATE_TABLE_SQL = '''
CREATE TABLE IF NOT EXISTS qa_knowledge (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    title TEXT,
    category TEXT,
    tag TEXT,
    content TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    UNIQUE(title, category, tag)
);
'''

def init_db():
    """Initialize the SQLite database and create the qa_knowledge table if it does not exist."""
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(CREATE_TABLE_SQL)
        conn.commit()

def upsert_qa(title: str, category: str, tag: str, content: str) -> None:
    """Insert or update a Q&A record. If title, category, tag match, update content and updated_at. Else insert new."""
    now = datetime.now(timezone.utc).isoformat()
    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.cursor()
        cur.execute("""
            SELECT id FROM qa_knowledge WHERE title=? AND category=? AND tag=?
        """, (title, category, tag))
        row = cur.fetchone()
        if row:
            cur.execute("""
                UPDATE qa_knowledge
                SET content=?, updated_at=?
                WHERE title=? AND category=? AND tag=?
            """, (content, now, title, category, tag))
        else:
            cur.execute("""
                INSERT INTO qa_knowledge (title, category, tag, content, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (title, category, tag, content, now, now))
        conn.commit()

def load_all_qa() -> List[Dict]:
    """Load all Q&A records as a list of dicts."""
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        cur = conn.cursor()
        cur.execute("SELECT * FROM qa_knowledge ORDER BY updated_at DESC, created_at DESC")
        rows = cur.fetchall()
        return [dict(row) for row in rows]

def update_qa_by_id(id: int, content: str) -> None:
    """Update the content and updated_at of a Q&A record by id."""
    now = datetime.now(timezone.utc).isoformat()
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            "UPDATE qa_knowledge SET content=?, updated_at=? WHERE id=?",
            (content, now, id)
        )
        conn.commit()

def delete_qa_by_id(id: int) -> None:
    """Delete a Q&A record by id."""
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute("DELETE FROM qa_knowledge WHERE id=?", (id,))
        conn.commit()


def is_version_table(headers):
    version_keywords = ['項番','パージョン', '改修日付', '改修者', '改修内容']
    return any(any(k in h for k in version_keywords) for h in headers)


def summarize_api_sections_with_nlp(sections):
    """
    セクションリストからパラメータ名・説明の対応やAPIリクエスト/レスポンス構造などを自然言語処理で整理する。
    現状は表や説明文から項目名と説明のマッピングを抽出する簡易版。
    戻り値: [{'section_title': ..., 'param_map': {項目名: 説明, ...}, 'raw_table': ...}, ...]
    """
    import re
    results = []
    for sec in sections:
        param_map = {}
        raw_table = None
        # パラメータ表がある場合は最初の表を使う
        if sec.get('tables'):
            table = sec['tables'][0]
            raw_table = table
            # 1行目がヘッダー（項目名）、2行目以降が値と仮定
            if len(table) > 1:
                headers = table[0]
                if is_version_table(headers):
                    continue
        # 本文から「項目名：説明」形式も抽出
        for line in sec.get('content', '').split('\n'):
            m = re.match(r'^(.+?)[：:](.+)$', line)
            if m:
                param, desc = m.group(1).strip(), m.group(2).strip()
                param_map[param] = desc
        results.append({
            'section_title': sec.get('title', ''),
            'param_map': param_map,
            'raw_table': raw_table
        })
    return results


def debug_word_structure(docx_file):
    """
    Wordファイルの各段落のテキストとスタイル名を出力して、構造をデバッグします。
    """
    from docx import Document
    doc = Document(docx_file)
    print("\n--- Word Document Structure Debug ---")
    for i, para in enumerate(doc.paragraphs):
        style_name = getattr(para.style, 'name', 'No Style')
        print(f"Para {i:03d}: Style='{style_name}', Text='{para.text.strip()}'")
    print("--- End of Debug ---\n")
    # Return an empty list to avoid breaking the app's flow
    return []

def extract_sections_by_toc_template(docx_file):
    """
    ハードコードされた目次テンプレートに基づき、Wordファイルから章・節・サブセクションを抽出します。
    最初の見出しが現れるまでのテキスト・テーブルは一切抽出しません。
    テキストとテーブルをドキュメントの出現順に処理し、正確に紐付けます。
    """
    toc_structure = [
        {"chapter": "一", "title": "インターフェース概要", "sections": [
            {"number": "1", "title": "通信プロトコルと伝送方式"},
            {"number": "2", "title": "パラメータ規定"},
            {"number": "3", "title": "デジタル署名"},
        ]},
        {"chapter": "二", "title": "インターフェース詳細", "sections": [
            {"number": "1", "title": "支払"},
            {"number": "2", "title": "返金"},
            {"number": "3", "title": "取消"},
            {"number": "4", "title": "オーダ照会"},
            {"number": "5", "title": "確認"},
        ]},
        {"chapter": "三", "title": "コード定義", "sections": [
            {"number": "1", "title": "通貨"},
            {"number": "2", "title": "取引ステータス"},
            {"number": "3", "title": "支払方法"},
            {"number": "4", "title": "オペレータシステム"},
            {"number": "5", "title": "エラーコード"},
        ]},
    ]

    all_headings = []
    for chap in toc_structure:
        for sec in chap["sections"]:
            pattern = rf"^{re.escape(sec['number'])}[.．]?\s*{re.escape(sec['title'])}"
            all_headings.append((pattern, chap['title'], sec['title']))
            if "subsections" in sec:
                for subsec in sec["subsections"]:
                    sub_pattern = rf"^{re.escape(subsec['number'])}[.．]?\s*{re.escape(subsec['title'])}"
                    all_headings.append((sub_pattern, chap['title'], subsec['title']))

    doc = Document(docx_file)
    sections_data = []
    current_section = None
    found_first_heading = False

    def flush_section():
        nonlocal current_section
        if current_section:
            current_section['content'] = current_section['content'].strip()
            if current_section['content'] or current_section['tables']:
                sections_data.append(current_section)
        current_section = None

    for block in doc.element.body:
        is_heading = False
        if isinstance(block, CT_P):
            para = Paragraph(block, doc)
            text = para.text.strip()
            if not text:
                continue

            for pattern, chapter_title, section_title in all_headings:
                if re.match(pattern, text):
                    found_first_heading = True
                    flush_section()
                    current_section = {
                        "chapter": chapter_title,
                        "section": section_title,
                        "title": text,
                        "content": '',
                        "tables": []
                    }
                    is_heading = True
                    break

            if not is_heading and found_first_heading:
                if not current_section:
                    flush_section()
                    continue
                current_section['content'] += text + '\n'

        elif isinstance(block, CT_Tbl):
            if not found_first_heading:
                continue  # 最初の見出し前のテーブルは無視
            if not current_section:
                flush_section()
                continue
            table = Table(block, doc)
            table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if table_data:
                current_section['tables'].append(table_data)

    flush_section()
    return sections_data
