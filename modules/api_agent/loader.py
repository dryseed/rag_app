from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph

def debug_word_structure(docx_file):
	"""
	Wordファイルの各段落のテキストとスタイル名を出力して、構造をデバッグします。
	"""
	doc = Document(docx_file)
	print("\n--- Word Document Structure Debug ---")
	for i, para in enumerate(doc.paragraphs):
		style_name = getattr(para.style, 'name', 'No Style')
		print(f"Para {i:03d}: Style='{style_name}', Text='{para.text.strip()}'")
	print("--- End of Debug ---\n")
	return []

def extract_sections_by_toc_template(docx_file):
	"""
	ハードコードされた目次テンプレートに基づき、Wordファイルから章・節・サブセクションを抽出します。
	"""
	toc_structure = [
		{"chapter": "一", "title": "インターフェース概要", "sections": [
			{"number": "1", "title": "通信プロトコルと伝送方式"},
			{"number": "2", "title": "パラメータ規定"},
			{"number": "3", "title": "デジタル署名"},
		]},
		{"chapter": "二", "title": "インターフェース詳細", "sections": [
			{"number": "1", "title": "支払"},
			{"number": "2", "title": "返金"},
			{"number": "3", "title": "取消"},
			{"number": "4", "title": "オーダ照会"},
			{"number": "5", "title": "確認"},
		]},
		{"chapter": "三", "title": "コード定義", "sections": [
			{"number": "1", "title": "通貨"},
			{"number": "2", "title": "取引ステータス"},
			{"number": "3", "title": "支払方法"},
			{"number": "4", "title": "オペレータシステム"},
			{"number": "5", "title": "エラーコード"},
		]},
	]

	all_headings = []
	for chap in toc_structure:
		for sec in chap["sections"]:
			pattern = rf"^{re.escape(sec['number'])}[.．]?\s*{re.escape(sec['title'])}"
			all_headings.append((pattern, chap['title'], sec['title']))
			if "subsections" in sec:
				for subsec in sec["subsections"]:
					sub_pattern = rf"^{re.escape(subsec['number'])}[.．]?\s*{re.escape(subsec['title'])}"
					all_headings.append((sub_pattern, chap['title'], subsec['title']))

	doc = Document(docx_file)
	sections_data = []
	current_section = None
	found_first_heading = False

	def flush_section():
		nonlocal current_section
		if current_section:
			current_section['content'] = current_section['content'].strip()
			if current_section['content'] or current_section['tables']:
				sections_data.append(current_section)
		current_section = None

	for block in doc.element.body:
		is_heading = False
		if isinstance(block, CT_P):
			para = Paragraph(block, doc)
			text = para.text.strip()
			if not text:
				continue

			for pattern, chapter_title, section_title in all_headings:
				if re.match(pattern, text):
					found_first_heading = True
					flush_section()
					current_section = {
						"chapter": chapter_title,
						"section": section_title,
						"title": text,
						"content": '',
						"tables": []
					}
					is_heading = True
					break

			if not is_heading and found_first_heading:
				if not current_section:
					flush_section()
					continue
				current_section['content'] += text + '\n'

		elif isinstance(block, CT_Tbl):
			if not found_first_heading:
				continue  # 最初の見出し前のテーブルは無視
			if not current_section:
				flush_section()
				continue
			table = Table(block, doc)
			table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
			if table_data:
				current_section['tables'].append(table_data)

	flush_section()
	return sections_data

def is_version_table(headers):
	version_keywords = ['項番','パージョン', '改修日付', '改修者', '改修内容']
	return any(any(k in h for k in version_keywords) for h in headers)

def summarize_api_sections_with_nlp(sections):
	"""
	セクションリストからパラメータ名・説明の対応やAPIリクエスト/レスポンス構造などを自然言語処理で整理する。
	"""
	import re
	results = []
	for sec in sections:
		param_map = {}
		raw_table = None
		if sec.get('tables'):
			table = sec['tables'][0]
			raw_table = table
			if len(table) > 1:
				headers = table[0]
				if is_version_table(headers):
					continue
		for line in sec.get('content', '').split('\n'):
			m = re.match(r'^(.+?)[：:](.+)$', line)
			if m:
				param, desc = m.group(1).strip(), m.group(2).strip()
				param_map[param] = desc
		results.append({
			'section_title': sec.get('title', ''),
			'param_map': param_map,
			'raw_table': raw_table
		})
	return results

import os
import re
import html
from typing import List, Dict
from modules.common import utils
from modules.common import embedding_model

def extract_markdown_sections_from_apidoc(base_dir="api_doc\onepay_cpm"):
	sections = []
	heading_pattern = r'^(#{1,4})\s+(.+)$'  # #, ##, ###, ####
	for root, dirs, files in os.walk(base_dir):
		spec_name = os.path.relpath(root, base_dir)
		if spec_name == ".":
			continue  # APIDoc直下はスキップ
		for file in files:
			if file.endswith(".md"):
				filepath = os.path.join(root, file)
				try:
					with open(filepath, "r", encoding="utf-8") as f:
						md_text = f.read()
				except Exception:
					continue
				md_text = re.sub(r'<[^>]+>', '', md_text)
				md_text = re.sub(r'InvalidCharacterError:.*', '', md_text)
				md_text = html.unescape(md_text)
				headings = [(m.start(), m.group(1), m.group(2).strip()) for m in re.finditer(heading_pattern, md_text, re.MULTILINE)]
				if not headings:
					content = md_text.strip()
					if not content or 'InvalidCharacterError' in content:
						continue
					sections.append({
						"content": content,
						"title": os.path.splitext(file)[0],
						"section": "全体",
						"spec_name": spec_name,
						"filename": file,
						"category": "API仕様書"
					})
					continue
				file_title = None
				for h in headings:
					if h[1] == '#':
						file_title = h[2]
						break
				if not file_title:
					file_title = os.path.splitext(file)[0]
				for idx, (start, level, title) in enumerate(headings):
					end = headings[idx+1][0] if idx+1 < len(headings) else len(md_text)
					content = md_text[start:end].strip()
					content_lines = content.splitlines()
					if content_lines and re.match(heading_pattern, content_lines[0]):
						content = '\n'.join(content_lines[1:]).strip()
					if not content or 'InvalidCharacterError' in content:
						continue
					sections.append({
						"content": content,
						"title": file_title,
						"section": title,
						"heading_level": len(level),
						"spec_name": spec_name,
						"filename": file,
						"category": "API仕様書"
					})
	return sections
